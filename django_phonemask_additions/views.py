from django.contrib.auth import authenticate, login, logout
from django.shortcuts import render, redirect
from django.utils import timezone
from django.contrib.auth.decorators import login_required, user_passes_test
from .forms import LoginForm, ChangePasswordForm, UserRegistrationForm, UserUpdateForm
from .models import CustomUser
import requests
import os
from docx import Document
from datetime import datetime

# Проверка админских прав
def check_admin(user):
    """Проверяет, является ли пользователь администратором."""
    return user.is_authenticated and user.role == 'admin'

# Получение данных от API
def fetch_api_data():
    """Извлекает данные от API симулятора."""
    try:
        response = requests.get('http://localhost:4444/TransferSimulator/mobilePhone', timeout=5)
        response.raise_for_status()
        return response.json()
    except requests.RequestException as e:
        return {'error': f"Ошибка API: {str(e)}"}

# Валидация данных API
def validate_api_data(data):
    """Проверяет, что данные от API содержат поле 'value'."""
    if not data:
        return False, "Данные от API не получены."
    if 'error' in data:
        return False, data['error']
    if 'value' not in data:
        return False, "Поле 'value' отсутствует в ответе API."
    return True, data['value']

# Сохранение данных в Word
def save_to_word(phone_data):
    """Сохраняет список номеров и результаты валидации в Word-файл."""
    try:
        doc = Document()
        doc.add_heading('Отчёт по валидации данных', 0)
        doc.add_paragraph(f'Дата и время: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        for data, result in phone_data:
            doc.add_paragraph(f'Номер телефона: {data}')
            doc.add_paragraph(f'Результат валидации: {result}')
            doc.add_paragraph('---')
        
        # Папка для сохранения
        os.makedirs('documents', exist_ok=True)
        filename = f'documents/report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        doc.save(filename)
        return filename
    except Exception as e:
        return f"Ошибка при создании документа: {str(e)}"

# Блокировка пользователя
def block_user(user, form, template='login.html'):
    """Блокирует аккаунт и возвращает ошибку."""
    user.is_blocked = True
    user.save()
    return render(None, template, {
        'form': form,
        'error': 'Аккаунт заблокирован. Свяжитесь с администратором.'
    })

# Успешный вход
def process_auth(user, request, password):
    """Обрабатывает успешную авторизацию."""
    user.failed_attempts = 0
    user.last_login_attempt = timezone.now()
    user.save()
    login(request, user)
    if user.role == 'admin' and password == 'initial_password':
        return redirect('password_change')
    return render(request, 'success.html', {
        'message': 'Вход выполнен успешно!'
    })

# Главная страница
def main_page(request):
    """Отображает главную страницу."""
    return render(request, 'home.html')

# Авторизация
def login_user(request):
    """Обрабатывает авторизацию пользователя."""
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']
            try:
                user = CustomUser.objects.get(username=username)
                if user.is_blocked or (
                    user.last_login_attempt and
                    (timezone.now() - user.last_login_attempt).days > 30
                ):
                    return block_user(user, form)
                auth_user = authenticate(request, username=username, password=password)
                if auth_user:
                    return process_auth(auth_user, request, password)
                user.failed_attempts += 1
                if user.failed_attempts >= 3:
                    return block_user(user, form)
                user.save()
                return render(request, 'login.html', {
                    'form': form,
                    'error': 'Неверный логин или пароль.'
                })
            except CustomUser.DoesNotExist:
                return render(request, 'login.html', {
                    'form': form,
                    'error': 'Пользователь не найден.'
                })
        return render(request, 'login.html', {'form': form})
    return render(request, 'login.html', {'form': LoginForm()})

# Выход
def logout_user(request):
    """Завершает сессию и перенаправляет на главную."""
    logout(request)
    return redirect('home_page')

@login_required
def update_password(request):
    """Обновляет пароль пользователя."""
    if request.method == 'POST':
        form = ChangePasswordForm(request.POST)
        if form.is_valid():
            user = request.user
            if user.check_password(form.cleaned_data['current_password']):
                new_pass = form.cleaned_data['new_password']
                confirm_pass = form.cleaned_data['confirm_password']
                if new_pass == confirm_pass:
                    user.set_password(new_pass)
                    user.save()
                    return render(request, 'success.html', {
                        'message': 'Пароль обновлён!'
                    })
                return render(request, 'change_password.html', {
                    'form': form,
                    'error': 'Новые пароли не совпадают.'
                })
            return render(request, 'change_password.html', {
                'form': form,
                'error': 'Текущий пароль неверный.'
            })
    return render(request, 'change_password.html', {
        'form': ChangePasswordForm()
    })

# Валидация данных API
def validate_data(request):
    """Обрабатывает валидацию данных от API и сохранение в Word."""
    bad_symbols = "=!@#$%^&*()_[]{}:;.<>?/|"
    
    # Инициализация списка номеров в сессии
    if 'phone_data' not in request.session:
        request.session['phone_data'] = []
    
    result = None
    data = None
    error = None
    saved_file = None

    if request.method == 'POST':
        if 'get_data' in request.POST:
            api_data = fetch_api_data()
            is_valid, message = validate_api_data(api_data)
            if is_valid:
                data = message
                result = "Найден запрещённый символ!" if any(c in bad_symbols for c in data) else "Запрещённые символы не найдены"
                request.session['phone_data'].append((data, result))
                request.session.modified = True
            else:
                error = message
        elif 'send_result' in request.POST:
            if request.session['phone_data']:
                saved_file = save_to_word(request.session['phone_data'])
                if "Ошибка при создании документа" in str(saved_file):
                    error = saved_file
                    saved_file = None
                else:
                    request.session['phone_data'] = []
                    request.session.modified = True
            else:
                error = "Нет данных для сохранения в документ."

    if request.session['phone_data']:
        data, result = request.session['phone_data'][-1]
    
    return render(request, 'validate.html', {
        'data': data,
        'result': result,
        'error': error,
        'saved_file': saved_file,
        'phone_data': request.session['phone_data']
    })

@login_required
@user_passes_test(check_admin)
def admin_control(request):
    """Панель управления для администраторов."""
    users = CustomUser.objects.all()
    return render(request, 'admin_panel.html', {
        'users': users
    })

@login_required
@user_passes_test(check_admin)
def register_user(request):
    """Создаёт нового пользователя."""
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data['username']
            if CustomUser.objects.filter(username=username).exists():
                return render(request, 'register_user.html', {
                    'form': form,
                    'error': 'Пользователь уже существует.'
                })
            CustomUser.objects.create_user(
                username=username,
                password=form.cleaned_data['password'],
                role=form.cleaned_data['role']
            )
            return redirect('admin_panel')
    return render(request, 'register_user.html', {
        'form': UserRegistrationForm()
    })

@login_required
@user_passes_test(check_admin)
def edit_user(request, user_id):
    """Редактирует данные пользователя."""
    user = CustomUser.objects.get(id=user_id)
    if request.method == 'POST':
        form = UserUpdateForm(request.POST, instance=user)
        if form.is_valid():
            form.save()
            return redirect('admin_panel')
    return render(request, 'update_user.html', {
        'form': UserUpdateForm(instance=user),
        'user': user
    })
